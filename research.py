import streamlit as st
import PyPDF2
import openai
from io import BytesIO
from docx import Document
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from collections import Counter
import re

# --- Streamlit Config ---
st.set_page_config(page_title="📑 Enhanced Research Paper Comparator", layout="centered")
st.title("📑 Research Paper Comparator with KPI & Visual Insights")

# --- OpenAI Key ---
api_key = st.text_input("Enter your OpenAI API Key", type="password")

# --- Upload PDFs ---
st.subheader("📤 Upload Two Research Papers (PDF)")
pdf1 = st.file_uploader("Upload First Research Paper", type=["pdf"], key="pdf1")
pdf2 = st.file_uploader("Upload Second Research Paper", type=["pdf"])

# --- Extract Text ---
def extract_text(uploaded_pdf):
    if not uploaded_pdf:
        return ""
    reader = PyPDF2.PdfReader(uploaded_pdf)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text[:4000]

# --- GenAI Prompt for KPI ---
def get_kpi_response(text1, text2):
    prompt = f"""
Compare these two research papers and answer ONLY in the following KPI format:

1. Domain Similarity (Yes/No + Explanation)
2. Research Aspect Overlap (Yes/No + What aspects are common)
3. Innovation Uniqueness Score (1-10 with justification)
4. Content Similarity Index (0-100% estimation)
5. Theme Summary (1-2 lines for each paper)
6. Gap Analysis (What one covers that other doesn't)
7. Best Use Case for Each Paper (brief description)

--- Paper 1 ---
{text1}

--- Paper 2 ---
{text2}
"""
    openai.api_key = api_key
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1000
    )
    return response.choices[0].message.content.strip()

# --- GenAI Prompt for Highlights ---
def get_highlights(text, paper_number):
    prompt = f"Give 3 important bullet-point highlights of the following research paper:\n\n{text}\n\n"
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=200
    )
    return response.choices[0].message.content.strip()

# --- Keyword Extraction & WordCloud ---
def extract_keywords(text, num_keywords=30):
    words = re.findall(r'\b[a-zA-Z]{5,}\b', text.lower())
    common = Counter(words).most_common(num_keywords)
    return dict(common)

def generate_wordcloud(keywords_dict, title):
    wordcloud = WordCloud(width=600, height=400, background_color='white').generate_from_frequencies(keywords_dict)
    fig, ax = plt.subplots()
    ax.imshow(wordcloud, interpolation='bilinear')
    ax.axis('off')
    ax.set_title(title, fontsize=16)
    st.pyplot(fig)

# --- Compare Button ---
if st.button("🔍 Compare and Generate Insights"):
    if not all([pdf1, pdf2, api_key]):
        st.error("Please upload both PDFs and enter your OpenAI API key.")
    else:
        with st.spinner("Analyzing and generating insights..."):
            try:
                t1 = extract_text(pdf1)
                t2 = extract_text(pdf2)

                # KPI Response
                kpi_result = get_kpi_response(t1, t2)

                # Highlights
                highlights1 = get_highlights(t1, 1)
                highlights2 = get_highlights(t2, 2)

                # Keyword clouds
                keywords1 = extract_keywords(t1)
                keywords2 = extract_keywords(t2)

                st.success("✅ Insights Generated!")

                # --- Display KPI Summary ---
                st.subheader("📊 KPI Summary")
                for line in kpi_result.split("\n"):
                    if ":" in line:
                        key, value = line.split(":", 1)
                        st.markdown(f"{key.strip()}: {value.strip()}")
                    else:
                        st.markdown(line.strip())

                # --- Visual Word Clouds ---
                st.subheader("☁ Keyword Cloud Comparison")
                col1, col2 = st.columns(2)
                with col1:
                    generate_wordcloud(keywords1, "Paper 1 Keywords")
                with col2:
                    generate_wordcloud(keywords2, "Paper 2 Keywords")

                # --- Key Highlights ---
                st.subheader("🔎 Key Highlights from Each Paper")
                st.markdown("Paper 1 Highlights:")
                st.markdown(highlights1)
                st.markdown("Paper 2 Highlights:")
                st.markdown(highlights2)

                # --- Export to DOCX ---
                doc = Document()
                doc.add_heading("Research Paper Comparison - KPI Report", 0)
                doc.add_paragraph(kpi_result)
                doc.add_heading("Paper 1 Highlights", level=1)
                doc.add_paragraph(highlights1)
                doc.add_heading("Paper 2 Highlights", level=1)
                doc.add_paragraph(highlights2)

                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                st.download_button("📥 Download KPI Report (DOCX)", data=buf,
                                   file_name="research_comparison_kpi.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            except Exception as e:
                st.error(f"Error: {e}")